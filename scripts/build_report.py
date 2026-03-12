"""Convert technical_report.md to a formatted .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH = "technical_report.md"
DOCX_PATH = "technical_report.docx"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Georgia", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline(para, text, bold=False, italic=False, code=False, size=11):
    run = para.add_run(text)
    font_name = "Courier New" if code else "Georgia"
    color = (0x44, 0x44, 0x44) if code else None
    set_font(run, name=font_name, size=10 if code else size, bold=bold, italic=italic, color=color)
    return run


def render_inline(para, text, base_size=11):
    """Parse inline markdown (bold, italic, inline code, links) and add runs."""
    # Pattern order matters: bold+italic, bold, italic, code, link
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"       # bold+italic
        r"|(\*\*(.+?)\*\*)"           # bold
        r"|(\*(.+?)\*)"               # italic
        r"|(`(.+?)`)"                 # inline code
        r"|(\[(.+?)\]\((.+?)\))"      # link
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            add_inline(para, text[last:m.start()], size=base_size)
        if m.group(1):
            add_inline(para, m.group(2), bold=True, italic=True, size=base_size)
        elif m.group(3):
            add_inline(para, m.group(4), bold=True, size=base_size)
        elif m.group(5):
            add_inline(para, m.group(6), italic=True, size=base_size)
        elif m.group(7):
            add_inline(para, m.group(8), code=True, size=base_size)
        elif m.group(9):
            # Render link text as underlined
            run = para.add_run(m.group(10))
            set_font(run, size=base_size)
            run.font.color.rgb = RGBColor(0x1a, 0x5f, 0xab)
            run.underline = True
        last = m.end()
    if last < len(text):
        add_inline(para, text[last:], size=base_size)


def add_table(doc, rows):
    """rows: list of lists of strings. First row is header."""
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            is_header = (i == 0)
            render_inline(p, cell_text.strip(), base_size=10)
            for run in p.runs:
                run.bold = is_header
                run.font.size = Pt(10)
            # Header row shading
            if is_header:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "E8E8E8")
                tcPr.append(shd)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)

    doc.add_paragraph()  # spacing after table


# ── Main builder ─────────────────────────────────────────────────────────────

def build_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        # ── Blank line ──
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule ──
        if stripped in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── H1 ──
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(p, before=0, after=4)
            render_inline(p, title_text, base_size=18)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            i += 1
            continue

        # ── H2 ──
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            para_spacing(p, before=14, after=4)
            run = p.add_run(stripped[3:])
            set_font(run, size=13, bold=True, color=(0x1a, 0x1a, 0x2e))
            i += 1
            continue

        # ── H3 ──
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            para_spacing(p, before=10, after=3)
            run = p.add_run(stripped[4:])
            set_font(run, size=11, bold=True, color=(0x33, 0x33, 0x33))
            i += 1
            continue

        # ── Bold line acting as label (e.g. **Stage 1: ...**) ──
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            para_spacing(p, before=6, after=2)
            render_inline(p, stripped, base_size=11)
            i += 1
            continue

        # ── Table ──
        if stripped.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                # Skip separator rows (---|--- etc)
                if re.match(r"^\|[\s\-\|:]+\|$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            add_table(doc, table_rows)
            continue

        # ── Bullet list ──
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            para_spacing(p, before=1, after=2)
            render_inline(p, stripped[2:], base_size=11)
            i += 1
            continue

        # ── Numbered list ──
        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25)
            para_spacing(p, before=1, after=2)
            render_inline(p, re.sub(r"^\d+\. ", "", stripped), base_size=11)
            i += 1
            continue

        # ── Italic-only line (figure captions, asides) ──
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            para_spacing(p, before=4, after=4)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            render_inline(p, stripped, base_size=10)
            i += 1
            continue

        # ── Regular paragraph ──
        p = doc.add_paragraph()
        para_spacing(p, before=2, after=6)
        p.paragraph_format.line_spacing = Pt(13.5)
        render_inline(p, stripped, base_size=11)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    build_docx(MD_PATH, DOCX_PATH)
